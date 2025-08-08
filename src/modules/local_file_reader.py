"""
Local File Reader Module
========================

This module handles reading and processing local files.
Supports PDF, DOCX, Excel, and CSV files.

Author: Report GPT Team
Date: 2024
"""

import os
from pathlib import Path
from typing import List, Dict, Optional
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try to import pandas with error handling
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
    logger.info("Pandas imported successfully")
except ImportError as e:
    logger.warning(f"Pandas not available: {e}")
    PANDAS_AVAILABLE = False


class LocalFileReader:
    """
    Local File Reader Class
    
    Handles reading and processing local files.
    Supports PDF, DOCX, Excel, and CSV files.
    """
    
    def __init__(self):
        """
        Initialize Local File Reader
        """
        self.supported_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv', '.txt']
        
    def read_pdf_file(self, file_path: str) -> Dict:
        """
        Read PDF file and extract text and tables
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            Dict: Extracted text and tables
        """
        try:
            import pdfplumber
            
            logger.info(f"Reading PDF file: {file_path}")
            
            text_content = []
            tables = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        text_content.append({
                            'page': page_num + 1,
                            'text': text
                        })
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table_num, table in enumerate(page_tables):
                        if table:
                            tables.append({
                                'page': page_num + 1,
                                'table_num': table_num + 1,
                                'data': table
                            })
            
            return {
                'text': text_content,
                'tables': tables,
                'file_path': file_path
            }
            
        except Exception as e:
            logger.error(f"Error reading PDF file: {e}")
            raise
    
    def read_docx_file(self, file_path: str) -> Dict:
        """
        Read DOCX file and extract text and tables
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            Dict: Extracted text and tables
        """
        try:
            from docx import Document
            
            logger.info(f"Reading DOCX file: {file_path}")
            
            doc = Document(file_path)
            text_content = []
            tables = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append({
                        'type': 'paragraph',
                        'text': para.text
                    })
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'table_num': table_num + 1,
                        'data': table_data
                    })
            
            return {
                'text': text_content,
                'tables': tables,
                'file_path': file_path
            }
            
        except Exception as e:
            logger.error(f"Error reading DOCX file: {e}")
            raise
    
    def read_doc_file(self, file_path: str) -> Dict:
        """
        Read DOC file and extract text and tables
        
        Args:
            file_path (str): Path to DOC file
            
        Returns:
            Dict: Extracted text and tables
        """
        try:
            # For .doc files, we'll use a simple text extraction approach
            # since .doc files are binary and require special handling
            logger.info(f"Reading DOC file: {file_path}")
            
            # Try to extract text using different methods
            text_content = []
            tables = []
            
            try:
                # Method 1: Try using Windows COM automation (requires Microsoft Word)
                import win32com.client
                import os
                
                # Create Word application object
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                try:
                    # Open the document
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    
                    # Extract text
                    text = doc.Content.Text
                    if text.strip():
                        text_content.append({
                            'type': 'extracted_text',
                            'text': text
                        })
                    
                    # Close document
                    doc.Close()
                    
                except Exception as e:
                    logger.warning(f"Windows COM method failed: {e}")
                finally:
                    # Quit Word application
                    word.Quit()
                
            except Exception as e:
                logger.warning(f"Could not use Windows COM: {e}")
                
                try:
                    # Method 2: Try using docx2txt for .doc files
                    import docx2txt
                    text = docx2txt.process(file_path)
                    if text.strip():
                        text_content.append({
                            'type': 'extracted_text',
                            'text': text
                        })
                    
                except Exception as e2:
                    logger.warning(f"Could not read .doc file with docx2txt: {e2}")
                
                try:
                    # Method 2: Try using python-docx (might work for some .doc files)
                    from docx import Document
                    doc = Document(file_path)
                    
                    # Extract text from paragraphs
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_content.append({
                                'type': 'paragraph',
                                'text': para.text
                            })
                    
                    # Extract tables
                    for table_num, table in enumerate(doc.tables):
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text for cell in row.cells]
                            table_data.append(row_data)
                        
                        if table_data:
                            tables.append({
                                'table_num': table_num + 1,
                                'data': table_data
                            })
                    
                except Exception as e2:
                    logger.warning(f"Could not read .doc file with python-docx: {e2}")
                
                # Method 2: Try using antiword (if available)
                try:
                    import subprocess
                    result = subprocess.run(['antiword', file_path], 
                                         capture_output=True, text=True)
                    if result.returncode == 0:
                        text_content.append({
                            'type': 'extracted_text',
                            'text': result.stdout
                        })
                    else:
                        logger.warning("antiword not available or failed")
                except Exception as e2:
                    logger.warning(f"antiword method failed: {e2}")
                
                # Method 3: Try using catdoc (if available)
                try:
                    import subprocess
                    result = subprocess.run(['catdoc', file_path], 
                                         capture_output=True, text=True)
                    if result.returncode == 0:
                        text_content.append({
                            'type': 'extracted_text',
                            'text': result.stdout
                        })
                    else:
                        logger.warning("catdoc not available or failed")
                except Exception as e3:
                    logger.warning(f"catdoc method failed: {e3}")
            
            return {
                'text': text_content,
                'tables': tables,
                'file_path': file_path
            }
            
        except Exception as e:
            logger.error(f"Error reading DOC file: {e}")
            return {
                'file_path': file_path,
                'error': f'Could not read DOC file: {e}',
                'text': [],
                'tables': []
            }
    
    def read_excel_file(self, file_path: str) -> Dict:
        """
        Read Excel file and extract data
        
        Args:
            file_path (str): Path to Excel file
            
        Returns:
            Dict: Extracted data from all sheets
        """
        if not PANDAS_AVAILABLE:
            raise ImportError("Pandas is required to read Excel files. Please install pandas: pip install pandas")
        
        try:
            logger.info(f"Reading Excel file: {file_path}")
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_data[sheet_name] = df.to_dict('records')
            
            return {
                'sheets': sheets_data,
                'file_path': file_path
            }
            
        except Exception as e:
            logger.error(f"Error reading Excel file: {e}")
            raise
    
    def read_csv_file(self, file_path: str) -> Dict:
        """
        Read CSV file and extract data
        
        Args:
            file_path (str): Path to CSV file
            
        Returns:
            Dict: Extracted data
        """
        if not PANDAS_AVAILABLE:
            raise ImportError("Pandas is required to read CSV files. Please install pandas: pip install pandas")
        
        try:
            logger.info(f"Reading CSV file: {file_path}")
            
            df = pd.read_csv(file_path)
            return {
                'data': df.to_dict('records'),
                'file_path': file_path
            }
            
        except Exception as e:
            logger.error(f"Error reading CSV file: {e}")
            raise
    
    def read_txt_file(self, file_path: str) -> Dict:
        """
        Read text file and extract content
        
        Args:
            file_path (str): Path to text file
            
        Returns:
            Dict: Extracted text content
        """
        try:
            logger.info(f"Reading text file: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                'text': [{'type': 'text', 'content': content}],
                'file_path': file_path
            }
            
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            raise
    
    def process_file(self, file_path: str) -> Dict:
        """
        Process file based on its extension or content
        
        Args:
            file_path (str): Path to file
            
        Returns:
            Dict: Processed file data
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        # If file has no extension, try to detect type from content
        if extension == '':
            try:
                # Try to read as PDF first (most common)
                return self.read_pdf_file(str(file_path))
            except Exception as e:
                logger.warning(f"Could not read as PDF, trying other formats: {e}")
                try:
                    # Try as DOCX
                    return self.read_docx_file(str(file_path))
                except Exception as e2:
                    logger.warning(f"Could not read as DOCX, trying Excel: {e2}")
                    try:
                        # Try as Excel
                        return self.read_excel_file(str(file_path))
                    except Exception as e3:
                        logger.warning(f"Could not read as Excel, trying CSV: {e3}")
                        try:
                            # Try as CSV
                            return self.read_csv_file(str(file_path))
                        except Exception as e4:
                            logger.warning(f"Could not read as CSV, trying text: {e4}")
                            try:
                                # Try as text file
                                return self.read_txt_file(str(file_path))
                            except Exception as e5:
                                logger.error(f"Could not read file in any format: {e5}")
                                return {'file_path': str(file_path), 'error': 'Could not detect file type'}
        
        # Process files with known extensions
        if extension == '.pdf':
            return self.read_pdf_file(str(file_path))
        elif extension == '.docx':
            return self.read_docx_file(str(file_path))
        elif extension == '.doc':
            return self.read_doc_file(str(file_path))
        elif extension in ['.xlsx', '.xls']:
            return self.read_excel_file(str(file_path))
        elif extension == '.csv':
            return self.read_csv_file(str(file_path))
        elif extension == '.txt':
            return self.read_txt_file(str(file_path))
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return {'file_path': str(file_path), 'error': 'Unsupported file type'}
    
    def process_directory(self, directory_path: str) -> List[Dict]:
        """
        Process all supported files in a directory
        
        Args:
            directory_path (str): Path to directory
            
        Returns:
            List[Dict]: List of processed file data
        """
        try:
            directory = Path(directory_path)
            if not directory.exists():
                raise ValueError(f"Directory not found: {directory_path}")
            
            processed_data = []
            
            for file_path in directory.iterdir():
                if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                    try:
                        data = self.process_file(str(file_path))
                        processed_data.append(data)
                        logger.info(f"Processed: {file_path.name}")
                    except Exception as e:
                        logger.error(f"Error processing {file_path.name}: {e}")
                        continue
            
            return processed_data
            
        except Exception as e:
            logger.error(f"Error processing directory: {e}")
            raise


# Example usage
if __name__ == "__main__":
    # Initialize reader
    reader = LocalFileReader()
    
    # Example: Process a single file
    # data = reader.process_file("path/to/your/file.pdf")
    
    # Example: Process a directory
    # data = reader.process_directory("data/input")
    
    print("Local File Reader initialized successfully!") 